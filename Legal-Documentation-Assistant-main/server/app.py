from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import os
import time
import logging
import sys
import PyPDF2
import docx
import requests
import MySQLdb
import re
import bcrypt
import mammoth
from document_embedding import collection, upsert_embeddings, query_by_text
from document_retrieval import retrieve_documents
from chatbot_model import EnhancedQAModel
from log_config import log_function_call, log_data_flow

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

# Load environment variables from .env file
load_dotenv()

# Database connection
try:
    db = MySQLdb.connect(
        database=os.getenv('DATABASE_NAME'),
        user=os.getenv('DATABASE_USER'),
        password=os.getenv('DATABASE_PASSWORD'),
        host=os.getenv('DATABASE_HOST'),
        port=int(os.getenv('DATABASE_PORT'))
    )
    logger.info("Database connection established successfully")
except Exception as e:
    logger.error(f"Database connection failed: {e}")
    db = None

# Create uploads directory if it doesn't exist
if not os.path.exists('./uploads'):
    os.makedirs('./uploads')
    logger.info("Created uploads directory")

# Create docs directory if it doesn't exist
if not os.path.exists('./docs'):
    os.makedirs('./docs')
    logger.info("Created docs directory")

@app.route('/api/register', methods=['POST'])
def register():
    data = request.get_json()
    name = data.get('name')
    phone = data.get('phone')
    email = data.get('email')
    password = data.get('password')
    
    # Validate input
    if not all([name, phone, email, password]):
        return jsonify({'message': 'All fields are required!'}), 400
    
    try:
        cur = db.cursor()
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        cur.execute("INSERT INTO users (name, phone, email, password) VALUES (%s, %s, %s, %s);", 
                    (name, phone, email, hashed_password))
        db.commit()
        cur.close()
        logger.info(f"User {email} registered successfully")
        return jsonify({'message': 'User registered successfully!'}), 201
    except MySQLdb.Error as e:
        logger.error(f"Error inserting user: {e}")
        return jsonify({'message': 'Error saving user to database!'}), 500

@app.route('/api/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')

    # Validate email and password
    if not email or not password:
        return jsonify({'message': 'Email and password are required!'}), 400

    try:
        # Check user credentials
        cur = db.cursor()
        cur.execute("SELECT id, name, password, logo FROM users WHERE email = %s;", (email,))
        user = cur.fetchone()
        cur.close()

        if user and bcrypt.checkpw(password.encode('utf-8'), user[2].encode('utf-8')):
            return jsonify({
                'message': 'Login successful!', 
                'user_id': user[0],
                'username': user[1], 
                'logo': user[3] if len(user) > 3 else None
            }), 200
        else:
            return jsonify({'message': 'Invalid email or password!'}), 401
    except Exception as e:
        logger.error(f"Login error: {e}")
        return jsonify({'message': 'Server error during login!'}), 500

@app.route('/api/upload', methods=['POST'])
def upload_document():
    if 'document' not in request.files:
        return jsonify({'message': 'No file part'}), 400
    
    file = request.files['document']
    if file.filename == '':
        return jsonify({'message': 'No selected file'}), 400
    
    try:
        extracted_text, error = process_uploaded_file(file)
        if error:
            return jsonify({'error': error}), 400
        
        # Generate a unique ID for the document
        unique_id = f"doc_{int(time.time())}"
        
        # Extract meaningful metadata
        file_extension = os.path.splitext(file.filename)[1].lower()
        metadata = {
            "source": file.filename,
            "text_length": len(extracted_text),
            "word_count": len(extracted_text.split()),
            "file_type": file_extension.replace('.', ''),
            "upload_timestamp": time.time(),
            "paragraphs": len([p for p in extracted_text.split('\n') if p.strip()])
        }
        
        # Insert document into ChromaDB collection with improved metadata
        try:
            # Use document_embedding function
            from document_embedding import upsert_embeddings
            
            upsert_success = upsert_embeddings(
                texts=[extracted_text],
                metadatas=[metadata],
                ids=[unique_id]
            )
            
            if not upsert_success:
                return jsonify({'error': 'Failed to store document in vector database.'}), 500
                
            logger.info(f"Document '{file.filename}' inserted with ID: {unique_id}")
            return jsonify({
                'message': 'File uploaded and processed successfully!',
                'document_id': unique_id,
                'file_name': file.filename,
                'word_count': metadata['word_count'],
                'paragraphs': metadata['paragraphs']
            }), 200
        except Exception as e:
            logger.error(f"Error inserting document into collection: {e}")
            return jsonify({'error': 'Failed to process document.'}), 500
    except Exception as e:
        logger.error(f"Error in document upload: {e}")
        return jsonify({'error': 'Server error during file upload'}), 500

def process_uploaded_file(file):
    """
    Process different file types and extract text content with enhanced validation
    
    Args:
        file: The uploaded file object
    
    Returns:
        tuple: (extracted_text, error_message)
    """
    try:
        file_path = os.path.join('./uploads', file.filename)
        file.save(file_path)
        logger.info(f"File saved to {file_path}")
        
        # Extract text based on file type
        text = ''
        extraction_method = ''
        
        if file.filename.lower().endswith('.pdf'):
            try:
                # First attempt with PyPDF2
                extraction_method = 'PyPDF2'
                reader = PyPDF2.PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text() or ''
                    if page_text.strip():
                        text += page_text + "\n\n"
                
                # Validate text extraction quality
                if not text.strip() or len(text) < 50:
                    logger.warning(f"Poor text extraction with PyPDF2 ({len(text)} chars). Trying fallback method.")
                    
                    # Try with pdfplumber as fallback
                    try:
                        import pdfplumber
                        extraction_method = 'pdfplumber'
                        text = ''  # Reset text
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                page_text = page.extract_text() or ''
                                text += page_text + "\n\n"
                    except ImportError:
                        logger.warning("pdfplumber not available for fallback PDF extraction")
            except Exception as pdf_error:
                logger.error(f"PDF extraction error: {pdf_error}")
                return None, f"Could not extract text from PDF: {str(pdf_error)}"
                
        elif file.filename.lower().endswith('.docx'):
            extraction_method = 'python-docx'
            doc = docx.Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Try mammoth for potentially better extraction if docx text is poor
            if not text.strip() or len(text) < 50:
                try:
                    extraction_method = 'mammoth'
                    with open(file_path, 'rb') as docx_file:
                        result = mammoth.extract_raw_text(docx_file)
                        text = result.value
                except Exception as mammoth_error:
                    logger.error(f"Mammoth extraction error: {mammoth_error}")
                    # Continue with python-docx result
        
        elif file.filename.lower().endswith('.txt'):
            extraction_method = 'text'
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
        else:
            return None, 'Unsupported file type. Please upload PDF, DOCX, or TXT files.'
        
        # Validate extraction results
        if not text.strip():
            return None, f'Could not extract text from the file using {extraction_method}.'
        
        # Text quality checks
        word_count = len(text.split())
        if word_count < 10:
            logger.warning(f"Very low word count ({word_count}) in extracted text")
            return None, f'Extracted text has too few words ({word_count}). The document may be scanned or corrupted.'
        
        # Check for common OCR/extraction issues
        weird_char_ratio = len(re.findall(r'[^\x00-\x7F]', text)) / max(len(text), 1)
        if weird_char_ratio > 0.3:  # If more than 30% are non-ASCII
            logger.warning(f"High ratio of non-ASCII characters: {weird_char_ratio:.2f}")
            # We still return the text but log the warning
        
        logger.info(f"Successfully extracted {len(text)} characters ({word_count} words) from {file.filename} using {extraction_method}")
        # Preview the first 200 chars for debugging
        logger.info(f"Text preview: {text[:200].replace('\n', ' ')}")
        
        return text, None
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {e}")
        return None, f"Error processing file: {str(e)}"
    
@app.route('/api/document-preview/<doc_id>', methods=['GET'])
def document_preview(doc_id):
    """
    Endpoint to preview extracted text from a document to verify what's stored in the vector database
    
    Args:
        doc_id: Document ID to preview
        
    Returns:
        JSON with document text and metadata
    """
    try:
        # Get document by ID from ChromaDB
        result = collection.get(ids=[doc_id])
        
        if not result or 'ids' not in result or not result['ids']:
            return jsonify({'error': 'Document not found'}), 404
        
        # Prepare text content and metadata
        doc_text = result['documents'][0] if 'documents' in result and result['documents'] else ""
        doc_metadata = result['metadatas'][0] if 'metadatas' in result and result['metadatas'] else {}
        
        # Text statistics for verification
        stats = {
            "total_chars": len(doc_text),
            "total_words": len(doc_text.split()),
            "total_paragraphs": len([p for p in doc_text.split('\n') if p.strip()]),
            "preview_length": min(500, len(doc_text))  # Limit preview to 500 chars
        }
        
        return jsonify({
            'document_id': doc_id,
            'source': doc_metadata.get('source', 'Unknown'),
            'text_preview': doc_text[:stats["preview_length"]],
            'stats': stats,
            'metadata': doc_metadata
        }), 200
    except Exception as e:
        logger.error(f"Error retrieving document preview for {doc_id}: {e}")
        return jsonify({'error': f'Error retrieving document preview: {str(e)}'}), 500

@app.route('/api/query', methods=['POST'])
@log_function_call

def query_document():
    data = request.get_json()
    if not data or 'query' not in data:
        return jsonify({'error': 'No query provided'}), 400
    
    query = data['query']
    logger.info(f"Query received: {query}")
    
    try:
        # Log query input
        log_data_flow('query_input', query, source='client', destination='qa_model')
        
        # Initialize the ChromaDB-enhanced QA model
        try:
            # Create model instance with proper ChromaDB configuration
            qa_model = EnhancedQAModel(
                model_name="deepset/roberta-base-squad2",
                chroma_collection_name="document_collection",  # Use consistent collection name
                chroma_db_path="./chroma_db"  # Use consistent ChromaDB path
            )
            
            # Process the query directly using ChromaDB retrieval
            result = qa_model.process_query(query, use_retrieved_docs=True)
            
            # Log the result
            logger.info(f"QA result: {result}")
            
            # Get source documents for attribution
            source_docs = []
            if 'source_documents' in result:
                for doc in result['source_documents']:
                    if isinstance(doc, dict) and 'metadata' in doc and 'source' in doc['metadata']:
                        source_docs.append(doc['metadata']['source'])
            
            # Build response
            response = result['answer']
            
            # Return the response with metadata
            return jsonify({
                'response': response,
                'confidence': float(result.get('confidence', 0.0)),
                'has_answer': result.get('has_answer', False),
                'sources': source_docs
            }), 200
        
        except Exception as model_error:
            logger.error(f"QA model error: {model_error}", exc_info=True)
            return jsonify({
                'response': f"Error processing your query: {str(model_error)}",
                'error': str(model_error)
            }), 200
    
    except Exception as e:
        logger.error(f"Error processing query: {e}", exc_info=True)
        return jsonify({
            'response': f"Server error: {str(e)}",
            'error': str(e)
        }), 200
    
@app.route('/api/document-text-preview', methods=['POST'])
@log_function_call
def document_text_preview():
    """
    Endpoint to verify extracted text before adding it to the database
    """
    if 'document' not in request.files:
        return jsonify({'message': 'No file part'}), 400
    
    file = request.files['document']
    if file.filename == '':
        return jsonify({'message': 'No selected file'}), 400
    
    try:
        # Use the improved document processing function
        extracted_text, error = process_uploaded_file(file)
        
        if error:
            return jsonify({'error': error}), 400
        
        # Text statistics for verification
        stats = {
            "total_chars": len(extracted_text),
            "total_words": len(extracted_text.split()),
            "total_paragraphs": len([p for p in extracted_text.split('\n') if p.strip()]),
            "preview_length": min(500, len(extracted_text))  # Limit preview to 500 chars
        }
        
        # Log the preview for debugging
        log_data_flow('document_preview', {
            'filename': file.filename,
            'stats': stats,
            'preview': extracted_text[:500]
        })
        
        return jsonify({
            'message': 'Text extracted successfully',
            'filename': file.filename,
            'text_preview': extracted_text[:stats["preview_length"]],
            'stats': stats
        }), 200
    except Exception as e:
        logger.error(f"Error in document preview: {e}")
        return jsonify({'error': f'Error generating preview: {str(e)}'}), 500
    
@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.get_json()
    if not data or 'input' not in data:
        return jsonify({'error': 'No input provided'}), 400
    
    query = data['input']
    logger.info(f"Chat input received: {query}")
    
    try:
        # Retrieve relevant documents
        retrieved_docs = retrieve_documents(query, top_k=3)
        
        if not retrieved_docs or 'matches' not in retrieved_docs or not retrieved_docs['matches']:
            return jsonify({'response': 'I don\'t have enough information to answer that question based on the uploaded documents.'}), 200
        
        # Initialize the Chatbot model
        chatbot = Chatbot()
        
        # Generate a response
        response = chatbot.generate_response(query, retrieved_docs['matches'])
        
        logger.info(f"Chat response generated for: '{query}'")
        return jsonify({'response': response}), 200
    except Exception as e:
        logger.error(f"Error in chat: {e}")
        return jsonify({'error': f'Error generating response: {str(e)}'}), 500

@app.route('/api/services', methods=["GET"])
def services():
    try:
        cur = db.cursor()
        cur.execute('SELECT * FROM services')
        row_headers = [x[0] for x in cur.description]
        rv = cur.fetchall()
        cur.close()
        
        json_data = []
        for result in rv:
            json_data.append(dict(zip(row_headers, result)))
        
        return jsonify(json_data), 200
    except Exception as e:
        logger.error(f"Error fetching services: {e}")
        return jsonify({'error': 'Error fetching services'}), 500

@app.route('/api/forms', methods=["GET"])
def get_forms():
    service_id = request.args.get('service_id')
    if not service_id:
        return jsonify({'error': 'Service ID is required'}), 400
    
    try:
        cur = db.cursor()
        cur.execute(
            "SELECT services.service_id, services.service_name, forms.form_id, forms.form_name, forms.form_link "
            "FROM services INNER JOIN forms ON services.service_id = forms.service_id "
            "WHERE forms.service_id = %s;", [service_id])
        
        row_headers = [x[0] for x in cur.description]
        rv = cur.fetchall()
        cur.close()
        
        json_data = []
        for result in rv:
            json_data.append(dict(zip(row_headers, result)))
        
        return jsonify(json_data), 200
    except Exception as e:
        logger.error(f"Error fetching forms: {e}")
        return jsonify({'error': 'Error fetching forms'}), 500

@app.route('/api/form-details', methods=["GET"])
def get_form_details():
    form_id = request.args.get('form_id')
    if not form_id:
        return jsonify({'error': 'Form ID is required'}), 400
    
    try:
        cur = db.cursor()
        
        # Get form details
        cur.execute("SELECT * FROM forms WHERE form_id = %s;", [form_id])
        form_headers = [x[0] for x in cur.description]
        form_data = cur.fetchall()
        
        # Get category details
        cur.execute(
            "SELECT * FROM ques_categories WHERE id IN "
            "(SELECT DISTINCT(category_id) FROM input_ques WHERE ques_id IN "
            "(SELECT form_query_id FROM form_queries WHERE form_id = %s));", [form_id])
        category_headers = [x[0] for x in cur.description]
        category_data = cur.fetchall()
        
        # Get question details
        cur.execute(
            "SELECT * FROM input_ques WHERE ques_id IN "
            "(SELECT form_query_id FROM form_queries WHERE form_id = %s);", [form_id])
        question_headers = [x[0] for x in cur.description]
        question_data = cur.fetchall()
        
        cur.close()
        
        # Combine all data
        json_data = []
        for result in form_data:
            json_data.append(dict(zip(form_headers, result)))
        
        for result in category_data:
            json_data.append(dict(zip(category_headers, result)))
        
        for result in question_data:
            json_data.append(dict(zip(question_headers, result)))
        
        return jsonify(json_data), 200
    except Exception as e:
        logger.error(f"Error fetching form details: {e}")
        return jsonify({'error': 'Error fetching form details'}), 500

@app.route('/api/final-content', methods=["POST"])
def final_content():
    form_details = request.json
    if not form_details or 'form_id' not in form_details:
        return jsonify({'error': 'Form ID is required'}), 400
    
    form_id = form_details["form_id"]
    
    try:
        # Get form link
        cur = db.cursor()
        cur.execute("SELECT form_link FROM forms where form_id = %s;", [form_id])
        result = cur.fetchone()
        cur.close()
        
        if not result:
            return jsonify({'error': 'Form not found'}), 404
        
        form_link = result[0]
        
        # Download the document
        response = requests.get(form_link)
        if response.status_code != 200:
            return jsonify({'error': 'Could not download form template'}), 500
        
        # Save the downloaded document
        file_path = './docs/localfile.docx'
        with open(file_path, 'wb') as f:
            f.write(response.content)
        
        # Process the document with replacements
        doc = docx.Document(file_path)
        
        # Get all numeric keys and sort them in reverse
        numeric_keys = [int(x) for x in form_details.keys() if x.isdigit()]
        numeric_keys.sort(reverse=True)
        
        # Replace placeholders with form values
        for key in numeric_keys:
            old_text = f'#{key}'
            new_text = str(form_details[str(key)])
            
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
        
        # Save the modified document
        output_path = "./docs/Output2.docx"
        doc.save(output_path)
        
        # Convert to HTML for preview
        with open(output_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
            html_content = result.value
        
        return jsonify({'content': html_content}), 200
    except Exception as e:
        logger.error(f"Error processing form content: {e}")
        return jsonify({'error': f'Error processing form: {str(e)}'}), 500

@app.route('/api/final-form', methods=["POST"])
def final_form():
    try:
        # Check if output file exists
        output_path = './docs/Output2.docx'
        if not os.path.exists(output_path):
            return jsonify({'error': 'Document not found'}), 404
    # Send the file as attachment
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        logger.error(f"Error sending final form: {e}")
        return jsonify({'error': f'Error downloading document: {str(e)}'}), 500

@app.route('/api/embed-documents', methods=['POST'])
def embed_documents_endpoint():
    data = request.get_json()
    if not data or 'documents' not in data:
        return jsonify({'error': 'No documents provided'}), 400
    
    documents = data['documents']
    
    try:
        # Process documents in batches
        max_retries = 3
        batch_size = 2  # Number of documents to insert in each batch
        
        for i in range(0, len(documents), batch_size):
            batch = documents[i:i + batch_size]
            
            for attempt in range(max_retries):
                try:
                    collection.add(
                        documents=[doc["document"] for doc in batch],
                        metadatas=[doc["metadata"] for doc in batch],
                        ids=[f"id{int(time.time())}_{j}" for j in range(i, i + len(batch))]  # Create unique IDs
                    )
                    logger.info(f"Batch of documents inserted successfully: {i} to {i + len(batch) - 1}")
                    break  # Exit retry loop on success
                except Exception as e:
                    logger.error(f"Attempt {attempt + 1} for batch {i} failed: {e}")
                    if attempt == max_retries - 1:
                        logger.error("All attempts to insert batch documents failed")
        
        return jsonify({'message': 'Documents embedded successfully!'}), 200
    except Exception as e:
        logger.error(f"Error embedding documents: {e}")
        return jsonify({'error': f'Error embedding documents: {str(e)}'}), 500

@app.route('/api/documents', methods=['GET'])
def list_documents():
    try:
        # Get all document IDs from the collection
        results = collection.get()
        
        if not results or 'ids' not in results or not results['ids']:
            return jsonify({'documents': []}), 200
        
        documents = []
        for i, doc_id in enumerate(results['ids']):
            metadata = results['metadatas'][i] if 'metadatas' in results else {}
            documents.append({
                'id': doc_id,
                'source': metadata.get('source', 'Unknown'),
                'timestamp': metadata.get('timestamp', None)
            })
        
        return jsonify({'documents': documents}), 200
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        return jsonify({'error': f'Error listing documents: {str(e)}'}), 500

@app.route('/api/document/<doc_id>', methods=['GET'])
def get_document(doc_id):
    try:
        # Get document by ID
        result = collection.get(ids=[doc_id])
        
        if not result or 'ids' not in result or not result['ids']:
            return jsonify({'error': 'Document not found'}), 404
        
        document = {
            'id': result['ids'][0],
            'text': result['documents'][0] if 'documents' in result else '',
            'metadata': result['metadatas'][0] if 'metadatas' in result else {}
        }
        
        return jsonify({'document': document}), 200
    except Exception as e:
        logger.error(f"Error retrieving document {doc_id}: {e}")
        return jsonify({'error': f'Error retrieving document: {str(e)}'}), 500

@app.route('/api/document/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    try:
        # Delete document by ID
        collection.delete(ids=[doc_id])
        logger.info(f"Document {doc_id} deleted successfully")
        return jsonify({'message': f'Document {doc_id} deleted successfully'}), 200
    except Exception as e:
        logger.error(f"Error deleting document {doc_id}: {e}")
        return jsonify({'error': f'Error deleting document: {str(e)}'}), 500
    
@app.route('/api/bot', methods=['POST'])
def bot():
    user_input = request.json
    
    # Choose one of the following models. get_response works using Bag of Words Principle while get_document works using Cosine Similarity
    # response = get_response(user_input['user_chat'])
    response = get_document(user_input['user_chat'])
    return jsonify({'aiMessage': response})

if __name__ == '__main__':
    logger.info("Starting the application...")
    app.run(debug=True, host='0.0.0.0', port=5000)